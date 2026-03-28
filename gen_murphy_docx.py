"""
Genera Murphy_Analisis_Tecnico_Resumen.docx en el escritorio.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── colores ──────────────────────────────────────────────────────────────────
BLUE  = RGBColor(0x3b, 0x82, 0xf6)
GREEN = RGBColor(0x10, 0xb9, 0x81)
RED   = RGBColor(0xef, 0x44, 0x44)
AMBER = RGBColor(0xf5, 0x9e, 0x0b)
GRAY  = RGBColor(0x64, 0x74, 0x8b)
WHITE = RGBColor(0xf1, 0xf5, 0xf9)
LGRAY = RGBColor(0x94, 0xa3, 0xb8)
DARK  = RGBColor(0x1e, 0x29, 0x3b)

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── helpers ──────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '3b82f6')
    pBdr.append(bot); pPr.append(pBdr)


def h2(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12); run.font.bold = True
    run.font.color.rgb = color


def h3(text, color=AMBER):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.bold = True
    run.font.color.rgb = color


def body(text, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color


def bullet(text, color=None, indent=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(indent * 0.4)
    run = p.add_run(f'•  {text}')
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color


def kv(key, val, kcolor=LGRAY, vcolor=WHITE):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(f'{key}: ')
    r1.font.size = Pt(10); r1.font.bold = True
    r1.font.color.rgb = kcolor
    r2 = p.add_run(val)
    r2.font.size = Pt(10)
    r2.font.color.rgb = vcolor


def pb():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
r = p.add_run('ANÁLISIS TÉCNICO DE LOS\nMERCADOS FINANCIEROS')
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('John J. Murphy')
r2.font.size = Pt(16); r2.font.color.rgb = LGRAY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Resumen Estructurado Completo — 19 Capítulos')
r3.font.size = Pt(12); r3.font.color.rgb = GRAY

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(
    'Filosofía · Dow Theory · Tendencias · S/R · Patrones de Precio\n'
    'Medias Móviles · Osciladores · Velas · Elliott · Ciclos\n'
    'Gestión del Dinero · Análisis Intermercado · Checklist')
r4.font.size = Pt(10); r4.font.color.rgb = GRAY; r4.font.italic = True

pb()

# ══════════════════════════════════════════════════════════════════════════════
# CAP 1 — FILOSOFÍA
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 1 — FILOSOFÍA DEL ANÁLISIS TÉCNICO')
body('El análisis técnico es el estudio de la acción del mercado (precio, volumen, interés abierto) '
     'mediante gráficos, con el propósito de pronosticar tendencias futuras de precios.')

h3('Los Tres Pilares')
kv('1. El mercado lo descuenta todo',
   'Todos los factores ya están reflejados en el precio. El analista técnico no necesita conocer las causas.')
kv('2. Los precios se mueven en tendencias',
   'Una tendencia en movimiento tiene mayor probabilidad de continuar que de revertir.')
kv('3. La historia se repite',
   'Los patrones reflejan la psicología humana, que es constante. Por eso los patrones funcionan.')

h3('Análisis Técnico vs. Fundamental')
bullet('Fundamental: estudia las CAUSAS del movimiento (balances, tasas, oferta/demanda).')
bullet('Técnico: estudia el EFECTO (el precio ya incorpora todo).')
bullet('El fundamental mueve el precio; el técnico lee ese precio.')
bullet('Ventaja: aplica a cualquier mercado/timeframe; el precio reacciona antes que los fundamentales.')

h3('Críticas al AT')
bullet('Profecía autocumplida → en realidad VALIDA la herramienta.')
bullet('Hipótesis del mercado eficiente (EMH) → la evidencia empírica muestra tendencias aprovechables.')
bullet('Subjetividad → real, pero la experiencia la reduce significativamente.')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 2 — DOW THEORY
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 2 — TEORÍA DE DOW')
body('Base del análisis técnico moderno. Charles Dow (1851-1902). Sistematizada por Hamilton y Rhea.')

h3('Seis Principios')
kv('1. Los promedios descuentan todo', 'Los índices reflejan toda la actividad e información del mercado.')
kv('2. Tres tipos de tendencia',
   'Primaria (1+ años) · Secundaria (3 sem-3 meses, retroceso 33-66%) · Menor (<3 semanas)')
kv('3. Tres fases del bull market',
   'Acumulación (informados compran) → Participación pública → Distribución (informados venden, euforia)')
kv('4. Confirmación entre índices',
   'Industrials Y Transportation deben confirmar la misma señal. Sin confirmación = señal inválida.')
kv('5. Volumen confirma la tendencia',
   'Vol sube con el precio en alcista; sube cuando cae en bajista.')
kv('6. La tendencia continúa hasta prueba de reversión',
   '"Si no se prueba lo contrario, la tendencia sigue." — El principio más importante.')

h3('Señales')
bullet('Alcista: máximo más alto → corrección → mínimo más alto → nuevo máximo. Confirmado por ambos índices.')
bullet('Bajista: mínimo más bajo → rebote → máximo más bajo → nuevo mínimo. Confirmado por ambos índices.')
body('Crítica: las señales llegan tarde (20-25% del movimiento ya ocurrió).', color=GRAY)

# ══════════════════════════════════════════════════════════════════════════════
# CAP 3-4 — GRÁFICOS Y TENDENCIA
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 3-4 — GRÁFICOS Y CONCEPTOS DE TENDENCIA')

h2('Tipos de Gráficos')
kv('Barras', 'Máximo, mínimo, apertura (izq.) y cierre (der.). El más usado en Occidente.')
kv('Líneas', 'Solo cierres. Más limpio, menos información.')
kv('Candlestick', 'Como barras pero con cuerpo coloreado. Verde = alcista, Rojo = bajista.')
kv('Punto y Figura', 'Sin eje temporal. Solo registra movimientos de precio significativos.')
kv('Escala Log vs. Aritmética', 'Log: distancias iguales = % iguales. Preferida para largo plazo.')

h2('Definición de Tendencia')
kv('Alcista', 'Máximos más altos Y mínimos más altos (higher highs + higher lows).')
kv('Bajista', 'Máximos más bajos Y mínimos más bajos (lower highs + lower lows).')
kv('Lateral', 'Sin dirección clara. También llamado rango o consolidación.')
kv('Temporal', 'Mayor (>1 año) · Intermedia (3 sem-6 meses) · Menor (<3 semanas)')

h2('Líneas de Tendencia')
bullet('Alcista: unir dos mínimos ascendentes. El tercer toque confirma validez.')
bullet('Bajista: unir dos máximos descendentes.')
bullet('Regla 1-3%: cerrar al menos 1-3% más allá para confirmar ruptura.')
bullet('Canal: línea paralela a la trendline; delimita el rango del movimiento.')
bullet('A más contactos y más tiempo vigente, más importante es la línea.')

h2('Soporte y Resistencia — Conceptos Clave')
kv('Soporte', 'Nivel donde el interés comprador detiene o revierte la caída.')
kv('Resistencia', 'Nivel donde el interés vendedor detiene o revierte el alza.')
body('★ INTERCAMBIO DE ROLES: Un soporte roto se convierte en resistencia y viceversa. '
     'Es uno de los conceptos más fiables del análisis técnico.', color=AMBER)
bullet('Números redondos (10, 50, 100, 1000): soporte/resistencia psicológicos.')
bullet('Importancia del nivel: depende del volumen negociado allí y del tiempo transcurrido.')

h2('Retrocesos (Retracements)')
bullet('Dow Theory: 33%, 50%, 66%. Más del 66-67% → posible reversión de tendencia.')
bullet('Fibonacci: 38.2%, 50%, 61.8%. El Golden Ratio (61.8%) = el retroceso más importante.')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 5 — PATRONES REVERSIÓN
# ══════════════════════════════════════════════════════════════════════════════
pb()
h1('CAP. 5 — PATRONES DE REVERSIÓN MAYORES')
body('Señalan que la tendencia previa está a punto de cambiar de dirección.')

h2('Hombro-Cabeza-Hombro (H&S) — El más fiable')
body('TECHO: Hombro izq. (vol alto) → Cabeza (nuevo máximo, vol menor) → Hombro der. (vol bajo).')
bullet('Neckline: línea que une los dos valles de la formación.')
bullet('Señal de venta: cierre POR DEBAJO de la neckline.')
bullet('Objetivo: distancia (cabeza → neckline) proyectada ABAJO desde el punto de ruptura.')
bullet('Pullback: el precio suele retroceder a la neckline antes de seguir. Segunda entrada.')
bullet('SUELO INVERTIDO: idéntico pero invertido. El volumen en el quiebre alcista es CRÍTICO.')

h2('Doble Techo ("M") y Doble Suelo ("W")')
bullet('Dos picos/valles en niveles similares, con corrección en el medio.')
bullet('IMPORTANTE: los dos picos deben estar separados por semanas o meses, no días.')
bullet('Confirmación: cierre más allá de la neckline (mínimo de la corrección).')
bullet('Objetivo: altura del patrón proyectada desde la neckline.')
bullet('Volumen: mayor en el primer pico, menor en el segundo.')

h2('Triple Techo / Suelo y Formaciones Redondeadas')
bullet('Triple: tres picos/valles similares. Si el central es mayor → es un H&S.')
bullet('Formación redondeada (Saucer): transición gradual. Confiable pero difícil en tiempo real.')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 6 — PATRONES CONTINUACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 6 — PATRONES DE CONTINUACIÓN')
body('Pausa en la tendencia seguida de continuación en la misma dirección.')

h2('Triángulos')
kv('Simétrico', 'Neutral — rompe generalmente en dirección de la tendencia previa. Esperar confirmación.')
kv('Ascendente', 'Resistencia horizontal + soporte ascendente. ALCISTA. Objetivo: altura del triángulo.')
kv('Descendente', 'Soporte horizontal + resistencia descendente. BAJISTA.')
body('Volumen: se contrae durante la formación, explota en la ruptura.')

h2('Banderas y Gallardetes — Los más fiables de continuación')
bullet('Se forman después de un movimiento brusco y vertical (el asta).')
bullet('Bandera: rectángulo inclinado contra la tendencia. Alcista: inclina levemente abajo.')
bullet('Gallardete: triángulo pequeño. Similar a la bandera pero convergente.')
bullet('Duración normal: 1-3 semanas. Más de 4 semanas pierde fiabilidad.')
body('★ "Las banderas vuelan a media asta." Objetivo: proyectar el largo del asta desde la ruptura.', color=AMBER)

h2('Cuñas, Rectángulo y Movimiento Medido')
kv('Cuña Ascendente', 'Líneas convergentes hacia arriba. Patrón BAJISTA. Rompe hacia abajo.')
kv('Cuña Descendente', 'Líneas convergentes hacia abajo. Patrón ALCISTA. Rompe hacia arriba.')
kv('Rectángulo', 'Oscila entre dos horizontales. Generalmente continuación, puede romper en cualquier dirección.')
kv('Movimiento Medido', 'Pierna A = Pierna C. El mercado sube, corrige, y sube otro tanto igual.')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 7 — VOLUMEN E INTERÉS ABIERTO
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 7 — VOLUMEN E INTERÉS ABIERTO')

h2('Principios del Volumen')
bullet('Vol sube + precio sube = tendencia alcista SANA.')
bullet('Vol baja + precio corrige = corrección NORMAL (saludable).')
bullet('Vol sube + precio baja = SEÑAL DE ALARMA en tendencia alcista.')
bullet('Vol baja + precio sube = debilidad, posible reversión.')
bullet('Ruptura + alto volumen = señal CONFIABLE.')
bullet('Ruptura + bajo volumen = sospechosa, probable falsa ruptura.')

h2('OBV (On Balance Volume) — Joe Granville')
body('Suma acumulativa: agrega volumen en días alcistas, resta en días bajistas.')
bullet('OBV hace nuevo máximo antes que el precio → anticipa ruptura alcista.')
bullet('OBV hace nuevo mínimo antes que el precio → anticipa ruptura bajista.')
bullet('Divergencia OBV/precio = señal de reversión potente.')

h2('Interés Abierto (OI) — Solo Futuros/Opciones')
body('Número de contratos vigentes. Sube al abrir posiciones; baja al cerrarlas.')

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Precio', 'Volumen', 'O.I.', 'Interpretación']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    p_ = cell.paragraphs[0]
    p_.clear()
    run = p_.add_run(h)
    run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE
    shade_cell(cell, '1e3a5f')

data = [
    ('↑ Sube', '↑ Sube', '↑ Sube', 'Mercado fuerte — tendencia alcista sana'),
    ('↑ Sube', '↓ Baja', '↓ Baja', 'Mercado débil — posible fin del alza'),
    ('↓ Baja', '↑ Sube', '↑ Sube', 'Mercado bajista fuerte'),
    ('↓ Baja', '↓ Baja', '↓ Baja', 'Mercado bajista debilitándose'),
]
for ri, row_data in enumerate(data):
    for ci, val in enumerate(row_data):
        cell = table.rows[ri + 1].cells[ci]
        p_ = cell.paragraphs[0]
        p_.clear()
        run = p_.add_run(val)
        run.font.size = Pt(9); run.font.color.rgb = WHITE
        shade_cell(cell, '111827' if ri % 2 == 0 else '0f1623')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# CAP 9 — MEDIAS MÓVILES
# ══════════════════════════════════════════════════════════════════════════════
pb()
h1('CAP. 9 — MEDIAS MÓVILES')

h2('Tipos')
kv('SMA (Simple)',
   'Suma de N cierres / N. Todos los pesos iguales. La más básica.')
kv('WMA (Ponderada lineal)',
   'Pesos 1, 2, …, N. El bar más reciente tiene peso N. Más reactiva que SMA.')
kv('EMA (Exponencial)',
   'k = 2/(N+1). EMA[hoy] = (Cierre × k) + (EMA_ayer × (1−k)). Incluye todo el historial. La más usada.')

h2('Parámetros Comunes')
bullet('20-21 días = aprox. 1 mes de trading (corto plazo).')
bullet('50 días = aprox. 2.5 meses (medio plazo).')
bullet('200 días = aprox. 10 meses, tendencia primaria (largo plazo).')

h2('Señales de Trading')
kv('Una MA', 'Precio sobre MA → alcista. Precio bajo MA → bajista.')
kv('Doble cruce', 'MA rápida sobre MA lenta → compra. Cruce inverso → venta.')
body('Golden Cross: SMA50 cruza sobre SMA200 → señal alcista de largo plazo.', color=GREEN)
body('Death Cross: SMA50 cruza bajo SMA200 → señal bajista de largo plazo.', color=RED)
kv('Triple cruce 4-9-18', 'MA4>MA9>MA18=alcista. MA4<MA9<MA18=bajista.')
body('★ Las MAs funcionan muy bien en tendencias y MUY MAL en mercados laterales (whipsaws).', color=AMBER)

h2('Bandas de Bollinger — John Bollinger')
kv('Central', 'SMA 20 períodos.')
kv('Superior', 'SMA20 + (2 × desviación estándar 20p).')
kv('Inferior', 'SMA20 − (2 × desviación estándar 20p). ~95% de precios dentro de las bandas.')
bullet('Squeeze (bandas se contraen): volatilidad baja → movimiento importante próximo.')
bullet('Precio en banda superior = sobrecomprado a corto plazo. Inferior = sobrevendido.')
bullet('En tendencia fuerte, el precio puede "caminar" por la banda superior/inferior.')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 10 — OSCILADORES
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 10 — OSCILADORES Y OPINIÓN CONTRARIA')
body('Miden el impulso (momentum). Más útiles en mercados laterales. '
     'Complementan la tendencia: NO la reemplazan.')
body('★ La DIVERGENCIA entre el oscilador y el precio es la señal más potente.', color=AMBER)

h2('RSI (Relative Strength Index) — J. Welles Wilder')
kv('Fórmula', 'RS = avg(ganancias N días) / avg(pérdidas N días). RSI = 100 − [100/(1+RS)].')
kv('Período', '14 días (recomendado). También 9 y 25. Rango: 0 a 100.')
bullet('>70 = sobrecomprado. <30 = sobrevendido. (En bull fuerte: usar 80/20).')
bullet('NO vender solo porque RSI=70. Es una alerta, no señal autónoma.')
bullet('Divergencia bajista: precio nuevo máximo, RSI no confirma → debilidad alcista.', color=RED)
bullet('Divergencia alcista: precio nuevo mínimo, RSI no confirma → posible rebote.', color=GREEN)
bullet('Nivel 50: RSI>50=alcista; RSI<50=bajista (contexto general).')
bullet('Rupturas de líneas de tendencia en el RSI anticipan las del precio.')

h2('MACD — Gerald Appel')
kv('Línea MACD', 'EMA12 − EMA26')
kv('Señal', 'EMA9 de la línea MACD')
kv('Histograma', 'MACD − Señal (visualiza la diferencia; cambia antes del cruce)')
bullet('MACD sobre señal → compra. MACD bajo señal → venta.')
bullet('Cruce de cero: más lento pero más confiable.')
bullet('Divergencia: la señal más potente del MACD.')
bullet('El histograma cambia de dirección antes que el cruce de líneas → anticipa la señal.')

h2('Estocástico (Stochastic) — George Lane')
kv('Fórmula', '%K = [(Cierre − Mínimo_N) / (Máximo_N − Mínimo_N)] × 100. %D = SMA3 del %K.')
kv('Período', 'N=14. Rango 0-100. Sobrecomprado >80, sobrevendido <20.')
bullet('%K cruza sobre %D desde zona sobrevendida → compra fuerte.')
bullet('Slow stochastic (suavizado extra) = más confiable, menos ruido.')

h2('ADX — J. Welles Wilder')
body('Mide la FUERZA de la tendencia, no su dirección. No tiene sesgo alcista/bajista.')
bullet('ADX > 25 y subiendo: tendencia fuerte → usar sistemas de seguimiento de tendencia.')
bullet('ADX < 20 y bajando: mercado lateral → usar osciladores de rango.')
bullet('+DI cruza sobre −DI (con ADX subiendo) → alcista. Inverso → bajista.')

h2('Opinión Contraria')
body('Cuando la gran mayoría tiene la misma opinión, el mercado se mueve en contra.')
bullet('>70-75% de alcistas en encuestas → señal contraria bajista.')
bullet('<30-35% de alcistas → señal contraria alcista.')
bullet('VIX > 40: pánico → oportunidad contraria de compra.', color=GREEN)
bullet('VIX < 12: complacencia → precaución, posible corrección.', color=RED)

# ══════════════════════════════════════════════════════════════════════════════
# CAP 12 — CANDLESTICKS
# ══════════════════════════════════════════════════════════════════════════════
pb()
h1('CAP. 12 — VELAS JAPONESAS (CANDLESTICKS)')
body('Desarrolladas por Honma en el siglo XVII. Introducidas en Occidente por Steve Nison.')
kv('Cuerpo blanco/verde', 'Cierre > Apertura. Alcista. Sombra superior = máximo; inferior = mínimo.')
kv('Cuerpo negro/rojo', 'Cierre < Apertura. Bajista.')

h2('Patrones de 1 Vela')
bullet('Marubozu blanco/verde: cuerpo grande sin sombras. Muy alcista.')
bullet('Marubozu negro/rojo: cuerpo grande sin sombras. Muy bajista.')
bullet('Doji: apertura ≈ cierre (indecisión). En tendencia larga → posible reversión.')
bullet('Martillo: cuerpo arriba, sombra inferior larga (>2x el cuerpo). En suelo bajista → alcista.')
bullet('Hombre colgado: igual que el martillo pero en TECHO → bajista.')
bullet('Estrella fugaz: sombra superior larga, cuerpo abajo. En techo → bajista.')

h2('Patrones de 2 Velas')
bullet('Engulfing alcista: vela roja + vela verde que envuelve completamente → alcista fuerte.', color=GREEN)
bullet('Engulfing bajista: vela verde + vela roja que envuelve completamente → bajista fuerte.', color=RED)
bullet('Harami: vela 2 dentro del cuerpo de vela 1 → desaceleración de la tendencia.')
bullet('Patrón penetrante: gap abajo + cierre sobre mitad del cuerpo previo → alcista en suelo.')
bullet('Nube oscura (Dark Cloud Cover): lo inverso al penetrante → bajista en techo.')

h2('Patrones de 3 Velas')
bullet('Morning Star: vela roja + pequeña con gap + vela verde. Señal de suelo. Muy alcista.', color=GREEN)
bullet('Evening Star: inverso del Morning Star. Señal de techo. Muy bajista.', color=RED)
bullet('Tres soldados blancos: tres velas verdes ascendentes. Alcista fuerte.', color=GREEN)
bullet('Tres cuervos negros: tres velas rojas descendentes. Bajista fuerte.', color=RED)
body('★ Más significativas en niveles S/R clave, con alto volumen y confirmadas por indicadores.', color=AMBER)

# ══════════════════════════════════════════════════════════════════════════════
# CAP 13 — ELLIOTT
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 13 — TEORÍA DE ONDAS DE ELLIOTT')
body('Ralph Nelson Elliott (años 30). Los mercados se mueven en patrones repetitivos de ondas '
     'que reflejan la psicología colectiva de los participantes.')

h2('Estructura 5 + 3 Ondas')
kv('Onda 1', 'Primer movimiento alcista. Débil, pocos lo reconocen.')
kv('Onda 2', 'Corrección. NUNCA retrocede más del 100% de onda 1.')
kv('Onda 3', 'La más larga y fuerte. NUNCA es la más corta. Gran volumen y participación.')
kv('Onda 4', 'Corrección. NO puede solapar con el territorio de precios de onda 1.')
kv('Onda 5', 'Último avance. Aparecen divergencias en osciladores. Euforia del mercado.')
kv('Onda A', 'Primera caída. Muchos la ven como corrección temporal.')
kv('Onda B', 'Rebote trampa para alcistas.')
kv('Onda C', 'Caída final. Generalmente igual o mayor que onda A.')

h2('Reglas Inquebrantables', color=RED)
bullet('Onda 2 NUNCA retrocede más del 100% de onda 1.', color=RED)
bullet('Onda 3 NUNCA es la más corta de las ondas impulsivas (1, 3, 5).', color=RED)
bullet('Onda 4 NUNCA solapa con el territorio de precio de onda 1.', color=RED)

h2('Fibonacci en Elliott')
bullet('Onda 2: retrocede ~61.8% de onda 1.')
bullet('Onda 4: retrocede ~38.2% de onda 3.')
bullet('Onda 3: se extiende al 161.8% de onda 1.')
bullet('Onda C: suele ser igual a onda A (ziz-zag) o 61.8% de A (corrección plana).')
body('Murphy: Elliott es muy subjetivo. Usarlo como marco de referencia, no como herramienta principal.', color=GRAY)

# ══════════════════════════════════════════════════════════════════════════════
# CAP 16 — GESTIÓN DEL DINERO
# ══════════════════════════════════════════════════════════════════════════════
pb()
h1('CAP. 16 — GESTIÓN DEL DINERO Y TÁCTICAS')
body('"Tres elementos del éxito: un método correcto, gestión del dinero y disciplina para aplicar ambos."')

h2('Reglas Fundamentales')
kv('Regla del 2%', 'Nunca arriesgar más del 2% del capital total en una sola operación.')
kv('Regla del 6%', 'Si las pérdidas del mes llegan al 6%, detener el trading ese mes.')
kv('R/R mínimo', '1:2 (arriesgar $1 para ganar $2). Ideal: 1:3 o más. Sin R/R favorable → no operar.')
kv('Position sizing', 'Unidades = (Capital × % riesgo) / (Entrada − Stop)')

h2('Tipos de Stop-Loss')
kv('Money Management Stop', 'Basado en el máximo $ a perder por operación.')
kv('Stop Técnico (PREFERIDO)', 'Basado en S/R. En largo: ligeramente por debajo del soporte clave.')
kv('Trailing Stop', 'Se mueve con el precio. Bloquea ganancias sin limitar el potencial.')
kv('Stop de Tiempo', 'Si en N días el precio no se movió como se esperaba, cerrar igual.')
body('★ NUNCA mover el stop para evitar ser sacado. Es el error más costoso del trading.', color=RED)

h2('Pirámide y Tácticas de Entrada')
bullet('Piramidear: agregar posiciones conforme el precio avanza a favor.')
bullet('Posición inicial = la más grande. Adiciones progresivamente menores.')
bullet('Comprar en correcciones (retrocesos al soporte) en lugar de perseguir el precio.')
bullet('Comprar en rupturas: usar stop-buy orders para entrada automática.')

h2('Errores Psicológicos a Evitar')
bullet('Promediar hacia abajo (comprar más en una posición perdedora).', color=RED)
bullet('Mover el stop para evitar ser sacado.', color=RED)
bullet('Operar emocionalmente (euforia, miedo, venganza).', color=RED)
bullet('Sobreoperar para recuperar pérdidas.', color=RED)
bullet('Dejar que una ganancia se convierta en pérdida.')
bullet('Operar contra la tendencia mayor.')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 17 — ANÁLISIS INTERMERCADO
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 17 — ANÁLISIS INTERMERCADO')
body('Los cuatro mercados principales están interrelacionados: Bonos · Acciones · Commodities · Dólar.')

h2('Relaciones Clave')
kv('Bonos ↔ Acciones', 'Correlación POSITIVA. Bonos lideran: suben primero, acciones los siguen.')
kv('Bonos ↔ Commodities', 'Correlación NEGATIVA. Commodities suben (inflación) → bonos bajan.')
kv('Dólar ↔ Commodities', 'Correlación NEGATIVA. Dólar fuerte → commodities caen (el oro especialmente).')

body('CADENA: Dólar sube → Commodities bajan → Inflación baja → Bonos suben → Acciones suben', color=GREEN)
body('CADENA: Dólar baja → Commodities suben → Inflación sube → Bonos bajan → Acciones caen', color=RED)

h2('Orden de Liderazgo en Recuperación')
bullet('1° Los bonos tocan suelo (tasas dejan de subir).')
bullet('2° Las acciones tocan suelo (meses después).')
bullet('3° Los commodities tocan suelo últimos (economía real más lenta).')

# ══════════════════════════════════════════════════════════════════════════════
# CAP 18 — INDICADORES DE MERCADO
# ══════════════════════════════════════════════════════════════════════════════
h1('CAP. 18 — INDICADORES DEL MERCADO DE ACCIONES')

h2('Amplitud (Breadth) — Línea A/D')
body('Suma acumulativa de (acciones que suben − acciones que bajan) por día.')
bullet('Divergencia bajista: índice nuevo máximo pero A/D no confirma → pocas acciones lideran → techo.', color=RED)
bullet('Divergencia alcista: índice nuevo mínimo pero A/D no confirma → suelo próximo.', color=GREEN)

h2('TRIN (Arms Index)')
body('= (Avances/Descensos) / (Volumen alzas / Volumen bajas)')
bullet('< 1: volumen en alzas. Alcista.  |  > 1: volumen en bajas. Bajista.')
bullet('TRIN > 2: pánico vendedor → posible suelo (contraria alcista).', color=GREEN)
bullet('TRIN < 0.5: euforia compradora → posible techo (contraria bajista).', color=RED)

h2('Sentiment: VIX y Put/Call')
bullet('Put/Call > 1.0-1.2: pesimismo → señal contraria alcista.', color=GREEN)
bullet('Put/Call < 0.5: optimismo excesivo → señal contraria bajista.', color=RED)
bullet('VIX > 40: miedo extremo → oportunidad de compra contraria.', color=GREEN)
bullet('VIX < 12-15: complacencia → precaución.', color=RED)

# ══════════════════════════════════════════════════════════════════════════════
# CHECKLIST + SEÑALES
# ══════════════════════════════════════════════════════════════════════════════
pb()
h1('CAP. 19 — CHECKLIST DE ANÁLISIS COMPLETO')

h2('1. Largo Plazo (Mensual/Semanal)')
bullet('¿Cuál es la tendencia primaria del mercado?')
bullet('¿Precio vs. MAs de largo plazo (50, 200)?')
bullet('¿Qué dice el análisis intermercado? (bonos, dólar, commodities)')

h2('2. Sectorial')
bullet('¿Cuáles sectores tienen mejor relative strength?')
bullet('¿El sector del activo es líder o rezagado?')

h2('3. Activo Individual — Las tres tendencias')
bullet('Largo plazo (mensual) · Mediano plazo (semanal) · Corto plazo (diario).')
bullet('¿Están las tres tendencias alineadas?')

h2('4. Soporte, Resistencia y Patrones')
bullet('¿Cuáles son los niveles S/R más importantes? ¿Hay patrón formado?')
bullet('¿El precio está cerca de un nivel crítico?')

h2('5. MAs, Volumen y Osciladores')
bullet('¿Precio sobre/bajo MAs clave (20, 50, 200)? ¿Golden Cross / Death Cross?')
bullet('¿El volumen confirma la tendencia?')
bullet('¿Hay divergencia en RSI, MACD o estocástico?')
bullet('¿El mercado está sobrecomprado o sobrevendido?')

h2('6. Gestión del Dinero')
bullet('¿Dónde va el stop-loss técnico?')
bullet('¿El R/R es mínimo 1:2?')
bullet('¿Cuántas unidades según la regla del 2%?')
bullet('¿Es un punto óptimo de entrada o espero corrección?')

doc.add_paragraph()

h1('SEÑALES — REFERENCIA RÁPIDA')

h2('Señales Alcistas', color=GREEN)
signals_buy = [
    'Ruptura por encima de resistencia con alto volumen',
    'Golden Cross (SMA50 > SMA200)',
    'RSI divergencia alcista (precio baja, RSI no)',
    'MACD cruza por encima de la señal',
    'H&S invertido con ruptura de neckline + volumen alto',
    'Doble suelo: ruptura de neckline confirmada',
    'Precio rebota en soporte con volumen alcista',
    'Martillo o envolvente alcista en soporte clave',
    'ADX sube y +DI > −DI',
    'OBV hace nuevo máximo antes que el precio',
    'Put/Call ratio extremadamente alto',
    'VIX > 40 (contrario)',
    'A/D Line divergencia alcista',
]
for s in signals_buy:
    bullet(s, color=GREEN)

h2('Señales Bajistas', color=RED)
signals_sell = [
    'Ruptura por debajo de soporte con alto volumen',
    'Death Cross (SMA50 < SMA200)',
    'RSI divergencia bajista (precio sube, RSI no)',
    'MACD cruza por debajo de la señal',
    'H&S con ruptura de neckline y volumen alto',
    'Doble techo: ruptura de neckline confirmada',
    'Precio rechazado en resistencia con volumen bajista',
    'Estrella fugaz o envolvente bajista en resistencia clave',
    'ADX sube y −DI > +DI',
    'OBV hace nuevo mínimo antes que el precio',
    'A/D Line divergencia bajista',
]
for s in signals_sell:
    bullet(s, color=RED)

h2('Filtros — Cómo Evitar Señales Falsas', color=AMBER)
bullet('Ruptura sin volumen = sospechosa. Esperar confirmación de al menos 1 día.', color=AMBER)
bullet('Confirmar cierre 1-3% más allá del nivel antes de actuar.', color=AMBER)
bullet('Buscar confluencia de múltiples señales (S/R + oscilador + vela + volumen).', color=AMBER)
bullet('Siempre operar en dirección de la tendencia mayor. Contra-tendencia = mayor riesgo.', color=AMBER)

pb()
h1('FIBONACCI — TABLA DE REFERENCIA')

h2('Retrocesos Clave')
kv('23.6%', 'Retroceso mínimo. Solo en tendencias muy fuertes.')
kv('38.2%', 'Corrección normal en tendencias fuertes. Nivel típico de onda 4 de Elliott.')
kv('50.0%', 'El retroceso más "psicológico" (Dow Theory). Muy respetado por el mercado.')
kv('61.8%', 'Golden Ratio (φ − 1). El retroceso más importante. Nivel típico de onda 2.')
kv('78.6%', 'Retroceso profundo. Si se supera suele implicar reversión de tendencia.')

h2('Extensiones Clave')
kv('100%',   'El precio recorre la misma distancia que el movimiento previo (onda C = onda A).')
kv('127.2%', 'Primer objetivo de extensión.')
kv('161.8%', 'Extensión clásica. Onda 3 = 1.618 × onda 1. La más común.')
kv('261.8%', 'Extensión de movimientos muy fuertes.')

h2('La Secuencia de Fibonacci')
body('1, 1, 2, 3, 5, 8, 13, 21, 34, 55, 89, 144, 233...')
body('Cada número / el siguiente → 0.618 (Golden Ratio). Cada número / el anterior → 1.618.')
body('1 − 0.618 = 0.382. Por eso los retrocesos del 61.8% y 38.2% son tan respetados.')

doc.add_paragraph()
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin = p_fin.add_run('— Fin del Resumen —')
r_fin.font.size = Pt(10); r_fin.font.italic = True; r_fin.font.color.rgb = GRAY

output = 'c:/Users/DELL/Desktop/Murphy_Analisis_Tecnico_Resumen.docx'
doc.save(output)

import os
size = os.path.getsize(output)
print(f'Guardado: {output}')
print(f'Tamaño: {size / 1024:.0f} KB')
